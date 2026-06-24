#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère le document Word de présentation de l'état d'avancement du projet Mon Prof Perso."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
A = lambda p: os.path.join(HERE, "assets", p)

GREEN = RGBColor(0x0E, 0x5A, 0x43)
ORANGE = RGBColor(0xE8, 0x72, 0x2A)
INK = RGBColor(0x18, 0x24, 0x1E)
MUTED = RGBColor(0x67, 0x75, 0x6D)

doc = Document()

# Styles de base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

def h(text, size=16, color=GREEN, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def para(text, size=11, color=INK, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.italic = italic
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.color.rgb = INK
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED

def phone_row(items):
    """items = list of (image_path, caption). Met les téléphones côte à côte dans un tableau."""
    items = [(p, c) for (p, c) in items if os.path.exists(p)]
    if not items: return
    n = len(items)
    t = doc.add_table(rows=2, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (img, cap) in enumerate(items):
        cell = t.cell(0, i)
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run().add_picture(img, width=Inches(1.85))
        cc = t.cell(1, i).paragraphs[0]; cc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cc.add_run(cap); rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def landscape(img, width=6.3, cap=None):
    if not os.path.exists(img): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Inches(width))
    if cap: caption(cap)

def table_two(rows, headers=("", "")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    hc = t.rows[0].cells
    for i, htext in enumerate(headers):
        r = hc[i].paragraphs[0].add_run(htext); r.bold = True; r.font.size = Pt(10)
    for a_, b_ in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(a_).font.size = Pt(10)
        c[1].paragraphs[0].add_run(b_).font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ====================================================================== PAGE DE TITRE
doc.add_paragraph().paragraph_format.space_before = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mon Prof Perso"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = GREEN
r2 = p.add_run("."); r2.bold = True; r2.font.size = Pt(40); r2.font.color.rgb = ORANGE
para("Soutien scolaire à domicile & en ligne — Côte d'Ivoire", size=15, color=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Document de présentation — État d'avancement du projet", size=13, color=INK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("23 juin 2026", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_before = Pt(20)
landscape(A("maquette_thumbnail.png"), width=3.2)
caption("Aperçu de la maquette d'origine (37 écrans)")

doc.add_page_break()

# ====================================================================== 1. LE PROJET
h("1. Le projet", 18)
para("Mon Prof Perso est une application mobile de mise en relation entre familles/élèves "
     "et professeurs particuliers vérifiés, pensée pour le marché ivoirien :")
bullet("professeurs vérifiés, cours à domicile ou en ligne (visio) ;")
bullet("paiement par Mobile Money (Orange Money, Wave, MTN MoMo) ;")
bullet("préparation aux examens (BEPC & BAC), cours en groupe ;")
bullet("suivi de progression, abonnement « prof attitré », parrainage.")
para("Le périmètre fonctionnel couvre trois rôles : parent, élève et professeur.")

h("Objectif de la mission", 14)
para("À partir de la maquette (37 écrans), produire deux applications 100 % natives "
     "— Android et iOS — fidèles au design, partageant la même API et la même base de données.")

# ====================================================================== 2. LA MAQUETTE
h("2. La maquette d'origine", 18)
para("La maquette décrit un parcours complet en 37 écrans, organisé en 11 étapes "
     "(démarrage & compte, recherche, réservation & paiement, apprentissage, suivi, "
     "espace professeur, abonnement, cours en groupe, gestion, reçus & réglages).")
landscape(A("maquette_A.png"), cap="Maquette — En-tête de marque & Étape A : Démarrage & compte")
landscape(A("maquette_B.png"), cap="Maquette — étapes suivantes (recherche, réservation…)")

doc.add_page_break()

# ====================================================================== 3. CE QUI A ÉTÉ RÉALISÉ
h("3. Ce qui a été réalisé", 18)
para("Deux applications natives séparées + un backend commun, le tout fonctionnel, "
     "compilé et testé sur émulateur/simulateur.")
table_two([
    ("Application Android", "Kotlin + Jetpack Compose — 37 écrans — build OK, lancée sur émulateur"),
    ("Application iOS", "Swift + SwiftUI — 37 écrans — build OK, lancée sur simulateur"),
    ("Backend commun", "Node.js / TypeScript + Express + PostgreSQL (Docker Compose)"),
    ("API partagée", "Mêmes endpoints & modèles consommés par les 2 apps"),
    ("Identité", "Marque « Mon Prof Perso », identifiant ci.monprofperso.app"),
], headers=("Composant", "Détail"))

h("Design system fidèle", 14)
para("Mêmes couleurs (vert #0E5A43, orange #E8722A, crème), mêmes polices Schibsted Grotesk "
     "(titres) et Hanken Grotesk (corps) bundlées dans les deux apps, icônes équivalentes "
     "(Material Icons côté Android, SF Symbols côté iOS).")

# ====================================================================== 4. CAPTURES DE L'APP
h("4. L'application en images", 18)

h("4.1 — Démarrage & choix du rôle (Android & iOS)", 13, color=INK)
phone_row([
    (A("app_welcome.png"), "Android — Bienvenue"),
    (A("app_ios_welcome.png"), "iOS — Bienvenue"),
    (A("app_signup.png"), "Inscription (rôle confirmé)"),
])
para("Le rôle est choisi une seule fois à l'accueil puis confirmé à l'inscription "
     "(plus de double saisie). Rendu identique sur les deux plateformes.", size=10, color=MUTED)

h("4.2 — Recherche & données en direct (API)", 13, color=INK)
phone_row([
    (A("app_home.png"), "Accueil"),
    (A("app_search.png"), "Résultats — API en direct"),
])
para("L'écran de résultats charge les professeurs depuis l'API commune (bandeau « Données "
     "en direct »), avec repli automatique sur les données locales si l'API est indisponible.",
     size=10, color=MUTED)

h("4.3 — Suivi de l'élève", 13, color=INK)
phone_row([
    (A("app_courses.png"), "Mes cours"),
    (A("app_progress.png"), "Suivi des progrès"),
])

h("4.4 — Espace professeur (parcours adapté au rôle)", 13, color=INK)
phone_row([
    (A("app_teacher_dashboard.png"), "Tableau de bord prof"),
    (A("app_teacher_account.png"), "Mon compte (prof)"),
])
para("Un professeur arrive directement sur son espace (revenus, demandes, agenda) et "
     "son écran « Mon compte » affiche son profil, et non celui d'un parent.",
     size=10, color=MUTED)

doc.add_page_break()

# ====================================================================== 5. ÉTAT D'AVANCEMENT
h("5. État d'avancement", 18)
table_two([
    ("Les 37 écrans (Android)", "Terminé"),
    ("Les 37 écrans (iOS)", "Terminé"),
    ("Backend API + base de données (Docker)", "Terminé"),
    ("Branchement live sur l'API (accueil, recherche, profil, cours, progrès)", "Terminé"),
    ("Navigation & boutons interactifs (sélecteurs, onglets, actions)", "Terminé"),
    ("Partage / copie natifs (reçu, parrainage)", "Terminé"),
    ("Parcours adapté au rôle (prof → espace prof)", "Terminé"),
    ("Rebrand complet « Mon Prof Perso »", "Terminé"),
    ("Authentification réelle (JWT, OTP SMS)", "À faire"),
    ("Paiement Mobile Money réel (sandbox)", "À faire"),
    ("Génération de PDF (reçus/bilans), carte, chat temps réel", "À faire"),
], headers=("Fonctionnalité", "Statut"))

h("Stack technique", 14)
bullet("Android : Kotlin 2.0, Jetpack Compose (Material 3), Navigation Compose, Retrofit.")
bullet("iOS : Swift 5.9, SwiftUI, NavigationStack, URLSession (async/await).")
bullet("Backend : Node.js / TypeScript, Express, PostgreSQL 16, Docker Compose.")
bullet("Ports (sans collision) : API 8099 · Postgres 5544 · Adminer 8098.")

h("6. Prochaines étapes", 18)
bullet("Authentification réelle (inscription/connexion, OTP SMS, sessions JWT).", )
bullet("Intégration paiement Mobile Money (CinetPay / Wave) en environnement de test.")
bullet("Fonctions avancées : génération PDF des reçus & bilans, carte des profs, messagerie temps réel.")
bullet("Tests automatisés et préparation de la mise en production (CI/CD, stores).")

doc.add_paragraph().paragraph_format.space_before = Pt(16)
para("Document généré automatiquement — Mon Prof Perso.", size=9, color=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER)

out = os.path.join(HERE, "Mon_Prof_Perso_Presentation.docx")
doc.save(out)
print("Document généré :", out)
